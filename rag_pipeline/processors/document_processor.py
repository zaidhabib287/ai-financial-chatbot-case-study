import hashlib
import os
import re
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document as DocxDocument
from pypdf import PdfReader


class DocumentProcessor:
    """Process various document formats and extract text"""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return ""

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"

            return text.strip()
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return ""

    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read().strip()
        except Exception as e:
            print(f"Error extracting text from TXT: {e}")
            return ""

    def extract_text(self, file_path: str) -> str:
        """Extract text from document based on file extension"""
        file_extension = Path(file_path).suffix.lower()

        if file_extension == ".pdf":
            return self.extract_text_from_pdf(file_path)
        elif file_extension == ".docx":
            return self.extract_text_from_docx(file_path)
        elif file_extension in [".txt", ".text"]:
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

    def clean_text(self, text: str) -> str:
        """Clean and preprocess text"""
        # Remove multiple spaces and newlines
        text = re.sub(r"\s+", " ", text)
        # Remove special characters but keep punctuation
        text = re.sub(r"[^\w\s\.\,\;\:\!\?\-\(\)]", "", text)
        # Strip leading/trailing whitespace
        text = text.strip()
        return text

    def chunk_text(self, text: str) -> List[Dict[str, Any]]:
        """Split text into chunks with overlap"""
        words = text.split()
        chunks = []

        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            chunk_words = words[i : i + self.chunk_size]
            chunk_text = " ".join(chunk_words)

            # Generate unique ID for chunk
            chunk_id = hashlib.md5(chunk_text.encode()).hexdigest()

            chunks.append(
                {
                    "id": chunk_id,
                    "text": chunk_text,
                    "start_index": i,
                    "end_index": min(i + self.chunk_size, len(words)),
                }
            )

        return chunks

    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Process a document and return structured data"""
        # Extract text
        raw_text = self.extract_text(file_path)

        # Clean text
        cleaned_text = self.clean_text(raw_text)

        # Create chunks
        chunks = self.chunk_text(cleaned_text)

        return {
            "file_path": file_path,
            "file_name": os.path.basename(file_path),
            "raw_text": raw_text,
            "cleaned_text": cleaned_text,
            "chunks": chunks,
            "num_chunks": len(chunks),
            "total_words": len(cleaned_text.split()),
        }


class RuleExtractor:
    """Extract compliance rules and sanctions from documents"""

    def __init__(self):
        self.rule_patterns = {
            "transfer_limit": r"(?:daily|per-day|maximum)\s+(?:transfer|limit)[\s\S]{0,50}(\d+(?:,\d+)*(?:\.\d+)?)\s*(?:BD|BHD|bahraini dinar)",
            "transaction_limit": r"(?:per|single|individual)\s+(?:transaction|transfer)[\s\S]{0,50}(\d+(?:,\d+)*(?:\.\d+)?)\s*(?:BD|BHD|bahraini dinar)",
            "blacklisted_country": r"(?:blacklisted|prohibited|sanctioned|restricted)\s+(?:countries|nations|jurisdictions)[\s\S]{0,200}([A-Z][a-zA-Z\s,]+)",
            "sanctions_name": r"(?:sanctioned|prohibited|blocked)\s+(?:individuals|entities|persons)[\s\S]{0,200}([A-Z][a-zA-Z\s,]+)",
        }

    def extract_rules(self, text: str) -> List[Dict[str, Any]]:
        """Extract compliance rules from text"""
        rules = []

        for rule_type, pattern in self.rule_patterns.items():
            matches = re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE)

            for match in matches:
                value = match.group(1).strip()

                # Clean numeric values
                if rule_type in ["transfer_limit", "transaction_limit"]:
                    value = re.sub(r"[,\s]", "", value)
                    try:
                        value = float(value)
                    except:
                        continue

                rules.append(
                    {
                        "rule_type": rule_type,
                        "rule_value": value,
                        "context": match.group(0),
                        "position": match.span(),
                    }
                )

        return rules

    def extract_sanctions_list(self, text: str) -> Dict[str, List[str]]:
        """Extract sanctions lists from text"""
        sanctions = {"countries": [], "entities": []}

        # Look for country lists
        country_pattern = r"(?:sanctioned|blacklisted|prohibited)\s+countries[:\s]+([A-Za-z,\s\-]+)(?:\.|;|\n)"
        country_matches = re.finditer(country_pattern, text, re.IGNORECASE)

        for match in country_matches:
            countries = match.group(1).split(",")
            sanctions["countries"].extend([c.strip() for c in countries if c.strip()])

        # Look for entity lists
        entity_pattern = r"(?:sanctioned|prohibited)\s+(?:entities|individuals|persons)[:\s]+([A-Za-z,\s\-]+)(?:\.|;|\n)"
        entity_matches = re.finditer(entity_pattern, text, re.IGNORECASE)

        for match in entity_matches:
            entities = match.group(1).split(",")
            sanctions["entities"].extend([e.strip() for e in entities if e.strip()])

        # Remove duplicates
        sanctions["countries"] = list(set(sanctions["countries"]))
        sanctions["entities"] = list(set(sanctions["entities"]))

        return sanctions
